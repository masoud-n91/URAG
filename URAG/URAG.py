# tools
import os
from datetime import datetime
import json

# RAG
from groq import Groq
from sentence_transformers import SentenceTransformer
from annoy import AnnoyIndex
import re

# documents
from docx import Document

# url
import requests
from bs4 import BeautifulSoup
from urllib.parse import urljoin
from urllib.request import urlopen, Request

 
class Chatter():
    def __init__(self, api_key: str, file_path: str= None, model_name: str = None, standalone_query_sys_prompt: str = None, final_response_sys_prompt: str = None):
    
        self.client = Groq(api_key=api_key,)
        
        if file_path:
            self.file_path = file_path
        else:
            self.file_path = None

        if model_name:
            self.model = SentenceTransformer(model_name)
        else:
            self.model = SentenceTransformer('all-MiniLM-L6-v2')


        if standalone_query_sys_prompt:
            self.standalone_query_sys_prompt = standalone_query_sys_prompt
        else:
            self.standalone_query_sys_prompt =  """
            Your role is to assist another chatbot by formulating standalone queries based on user-submitted Queries and any relevant chat history.

            Guidelines:
            Create Standalone Queries: Only if the current query is unclear and related to the chat history, formulate a query that can be understood independently of the chat history, otherwise do not change the query.
            Handling Chat History: Use the chat history to inform the phrasing of your query. If no chat history is present or the current query is not related to the chat history, use the query as given.
            Exclusions: Do not include the name of the master or the name of the university in the query unless it is essential for context.
            Query Output: Provide only the standalone query; do not attempt to answer the question yourself.
            Clear Response: Do not add any introductory phrases, or any text other than the query itself (e.g., avoid phrases like "Here is your standalone query").
            Personal References: Replace any instance of "you" in the query with "chatbot".
            """

        if final_response_sys_prompt:
            self.faq_sys_prompt = final_response_sys_prompt
        else:
            self.final_response_sys_prompt =  """
            You are a friendly assistant named Laura, representing a program at a university.
            If the query is presented in a specific language, your response must also be in that same language to maintain consistency.
            If you cannot find the suitable answer for the user's query, you should respond with "Sorry, I don't have an answer to that question. Please try again with a different question."
            If a query is not clear for you, ask for a revised question.

            Instruction:
            Based on the given the top responses, chat history and the user's query, generate a response.
            The user content will be provided in the following format:
            - Query: a string representing the user's query
            - Chat history: a string representing the chat history, with each message separated by a newline character
            - Top responses: a string representing the top responses, with each response separated by a newline character
            Avoid giving anything rather than the response. DO NOT USE pharases like "Here is the response:" or "Here is a response".

            Query: here you will get a query from user
            """


    def read_docx(self):
        if self.filepath:
            file_path = self.file_path
        else:
            # file_path is the current working directory
            file_path = os.getcwd() + "data/"

        try: 
            doc = Document(file_path)
            contents = []
            temp = ""
            for paragraph in doc.paragraphs:
                if len(paragraph.text) > 0:
                    if paragraph.text[0] != "#":
                        temp += paragraph.text + "\n"
                    else:
                        contents.append(temp)
                        temp = ""
        except:
            print("File not found")
            contents = ""
        return contents


    def preprocess_text(self, text):
        text = text.lower()
        text = re.sub(r'\W+', ' ', text)
        return text


    def create_standalone_query(self, query: str, chat_history: str):
        chat_completion = self.client.chat.completions.create(
                messages=[
                    {
                        "role": "system",
                        "content": self.standalone_query_sys_prompt,
                    },
                    {
                        "role": "user",
                        "content": "Here is the user's query:\n{" + query + "} End of the user's query.\n\nHere is the chat history:\n{" + chat_history + "} End of the chat history.",
                    },
                ],
                model="llama3-8b-8192",
                temperature=0.5,
                top_p=1,
                stop=None,
                stream=False,
            )
        
        return chat_completion.choices[0].message.content


    def generate_final_response(self,top_responses: list, query: str, chat_history: str):
        now = datetime.now().strftime("%A, %B %d, %Y")

        sys_prompt = f"Today is {now}. " + self.final_response_sys_prompt

        top_strs = ""
        for idx, response in enumerate(top_responses):
            top_strs += f"{idx}- {response}\n\n"

        for i in range(10):
            try:
                chat_completion = self.client.chat.completions.create(
                        messages=[
                            {
                                "role": "system",
                                "content": sys_prompt,
                            },
                            {
                                "role": "user",
                                "content": "Here is the user's query:\n{" + query + "} End of the user's query.\n\nHere is the chat history:\n{" + chat_history + "} End of the chat history.\n\nHere are the top responses:\n{" + top_strs + "} End of the top responses.",
                            },
                        ],
                        model="llama3-70b-8192",
                        temperature=1.7,
                        top_p=1,
                        stop=None,
                        stream=False,
                    )
                
                json_data = {"text": chat_completion.choices[0].message.content}

                break

            except Exception as e:
                print("#################################################################################")
                print(e)

        final_json = self.move_urls_from_text(json_data)

        return final_json


    def move_urls_from_text(self, json_data):
        urls = re.findall(r'https?://\S+', json_data["text"])
        count = 0
            
        def replace_url(match):
            nonlocal count
            if match.group(0).endswith((',', '.')):
                ending = match.group(0)[-1]
            else:
                ending = ""
            count += 1
            return f'{"first" if count == 1 else "second" if count == 2 else "third" if count == 3 else {count}+"th"} link{ending}'

        modified_text = re.sub(r'https?://\S+', replace_url, json_data["text"])

        urls = [url[:-1] if url.endswith((',', '.')) else url for url in urls]
        
        json_data["website"] = urls
        json_data["text"] = modified_text
        
        return json_data


    def save_index(self, index, path, dimension: int):
        index.save(path +str(dimension) + ".ann")


    def create_vector_database(self, embeddings, path, n_trees=10):
        dimension = embeddings[0].shape[0]  
        annoy_index = AnnoyIndex(dimension, 'angular')  
        
        for i, vector in enumerate(embeddings):
            annoy_index.add_item(i, vector)
        
        annoy_index.build(n_trees)
        self.save_index(annoy_index, path, dimension)
        return annoy_index

    def find_documents(self, directory: str):
        document_embeddings = []
        documents = []

        for filename in os.listdir(directory):
            if filename.endswith('.docx'):
                file_path = os.path.join(directory, filename)
                contents = self.read_docx(file_path)
                for content in contents:
                    processed_content = self.preprocess_text(content)
                    embedding = self.model.encode(processed_content)
                    document_embeddings.append(embedding)
                    documents.append(content)

        return documents, document_embeddings


    def find_ann_files(self, directory):
        for filename in os.listdir(directory):
            if filename.endswith('.ann'):
                return filename


    def load_index(self, filename, dimension):
        loaded_index = AnnoyIndex(dimension, 'angular')
        loaded_index.load(filename)
        return loaded_index


    def get_similar_documents(self, query_vector, index, top_k=5):
        num_items = index.get_n_items()
        if top_k > num_items:
            top_k = num_items
        indices, distances = index.get_nns_by_vector(query_vector, top_k, include_distances=True)
        return indices, distances


    def extract_all_endpoints(self, url):
        url_len = len(url)
        endpoints = self.get_links(url)
        urls = []

        for endpoint in endpoints:
            if endpoint[:url_len]==url:
                urls.append(endpoint)

        return urls


    def get_links(self, url):
        links = set()
        try:
            response = requests.get(url)
            soup = BeautifulSoup(response.text, 'html.parser')
            for link in soup.find_all('a'):
                href = link.get('href')
                if href:
                    full_link = urljoin(url, href)
                    links.add(full_link)
        except requests.exceptions.RequestException as e:
            print(f"An error occurred: {e}")
        return links


    def write_urls_to_docx(self, urls, docx_filename):
        doc = Document()
        
        for url in urls:
            try:
                req = Request(url, headers={'User-Agent': 'Mozilla/5.0'})
                response = urlopen(req)
                encoding = response.headers.get_content_charset('utf-8')
                html = response.read().decode(encoding)
            except Exception as e:
                print(f"Failed to read {url}: {str(e)}")
                continue

            header = "# " + url
            doc.add_paragraph(header)
            soup = BeautifulSoup(html, features="html.parser")

            for script in soup(["script", "style"]):
                script.decompose() 

            text = soup.get_text()

            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)

            text = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', text)
            text = re.sub(r'[^\x00-\x7F]+', ' ', text)  

            doc.add_paragraph(text)

        doc.save(docx_filename)


    def read_url(self, directory, url:str = ""):
        filename = directory + "url - " + url.split("/")[-2] + ".docx"
        urls = self.extract_all_endpoints(url)
        self.write_urls_to_docx(urls, filename)


    def download_docx(self, docx_link, directory):
        response = requests.get(docx_link)
        response.raise_for_status()

        file_path = self.get_new_faq_file_name(directory)

        with open(file_path, 'wb') as f:
            f.write(response.content)


    def get_new_faq_file_name(self, directory):
        if not os.path.exists(directory):
            return None
        
        all_files = os.listdir(directory)

        faq_files = [file for file in all_files if file.startswith("FAQ")]
        file_path = directory + "FAQ" + str(len(faq_files) + 1) + ".docx"

        return file_path

class Memory():
    def __init__(self):
        self.chat_history = {}
        self.chat_history5 = {}

    
    def remove_mem(self, session_id):
        del self.chat_history[session_id]
        del self.chat_history5[session_id]


    def get_last_mem(self, session_id):
        if session_id not in self.chat_history:
            self.chat_history[session_id] = ""

        str_mem = ""

        str_mem = self.chat_history[session_id]

        if str_mem == "":
            str_mem = "No chat history yet"
    
        return str_mem
    

    def get_5_mem(self, session_id):
        if session_id not in self.chat_history5:
            self.chat_history5[session_id] = []

        str_mem = ""

        for i, mem in enumerate(self.chat_history5[session_id]):
            str_mem += f"{i}- {mem}\n\n"

        if str_mem == "":
            str_mem = "No chat history yet"
    
        return str_mem
    

    def update_chat(self, session_id, user_content, final_response):
        
        if len(self.chat_history5[session_id]) < 5:
            self.chat_history5[session_id].append("USER:\n" + user_content + "\n\nASSISTANT:\n" + str(final_response))
        else:
            self.chat_history5[session_id].pop(0)
            self.chat_history5[session_id].append("USER:\n" + user_content + "\n\nASSISTANT:\n" + str(final_response))

        self.chat_history[session_id] = "USER:\n" + user_content + "\n\nASSISTANT:\n" + str(final_response)